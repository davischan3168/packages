#!/usr/bin/env python3
# -*-coding:utf-8-*-

import os
from os.path import basename
import sys
import re
#import pickle
#import codecs
#import string
#import shutil
from docx import Document
import subprocess
from zipfile import ZipFile
#need to install python-docx
try:
    import win32com
    from win32com.client import Dispatch, constants
except:
    pass

def _read_doc(path):
    """
    读取doc文件,path是doc文件的路径。
    
    temp='text.txt'
    path=os.path.abspath(path)
    if sys.platform=='win32':
        word=Dispatch('Word.Application')
        doc = word.Documents.Open(path)
        doc.SaveAs(temp, 4)
        doc.Close()

        f = open(temp,'r')
        content= f.read()
        f.close()
        
        os.remove(temp)
        
        content = content.replace(" ","")
        #content=re.findall(r"[\u4e00-\u9fa5]+",content)
        #content='\n'.join(content)
    """
    path=os.path.abspath(path)
    if sys.platform=='win32':
        
        word=Dispatch('Word.Application')
        word.Visible = 0
        doc = word.Documents.Open(path)
        fullText=[]
        paras=doc.paragraphs
        for p in paras:
            fullText.append(p.Range.Text)
        content='\n'.join(fullText)          
        return content
    
    elif sys.platform=='linux':
        try:
            content = os.popen('catdoc %s'%path)
            #need to install catdoc for linux,for example:aptitude install catdoc
            content = content.read()
        except:
            content = subprocess.check_output(['antiword',path])
        return content

def _read_docx(path):
    """
    读取docx文件,path为docx文件的路径.
    有关文件可以阅读：http://www.cnblogs.com/wrajj/p/4914102.html
    """
    path=os.path.abspath(path)
    doc=Document(path)
    fullText=[]
    paras=doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    content='\n'.join(fullText)
    return content

def Doc2Docx(path):
    
    path=os.path.abspath(path)
    name=path+'x'
    if sys.platform=='win32':
        word=Dispatch('Word.Application')
        word.Visible = 0
        doc = word.Documents.Open(path)
        
        if os.path.exists(name):
            os.remove(name)
        doc.SaveAs(name,12, False, "", True, "", False, False, False, False)
        doc.Close()
        #word.Quit
    elif sys.platform=='linux':
        output = subprocess.check_output(["soffice","--headless","--invisible","--convert-to","docx",path,"--outdir",os.path.split(path)[0]])

    return
########################################
def MSOffice_change_type(docpath,dest_ftype='docx'):
    """
    wdFormatDocument = 0
    wdFormatDocument97 = 0
    wdFormatDocumentDefault = 16
    wdFormatDOSText = 4
    wdFormatDOSTextLineBreaks = 5
    wdFormatEncodedText = 7
    wdFormatFilteredHTML = 10
    wdFormatFlatXML = 19
    wdFormatFlatXMLMacroEnabled = 20
    wdFormatFlatXMLTemplate = 21
    wdFormatFlatXMLTemplateMacroEnabled = 22
    wdFormatHTML = 8
    wdFormatPDF = 17
    wdFormatRTF = 6
    wdFormatTemplate = 1
    wdFormatTemplate97 = 1
    wdFormatText = 2
    wdFormatTextLineBreaks = 3
    wdFormatUnicodeText = 7
    wdFormatWebArchive = 9
    wdFormatXML = 11
    wdFormatXMLDocument = 12
    wdFormatXMLDocumentMacroEnabled = 13
    wdFormatXMLTemplate = 14
    wdFormatXMLTemplateMacroEnabled = 15
    wdFormatXPS = 18
    """
    ftype={'doc':0,'docx':16,'untxt':7,'html':8,'txt':2,'pdf':17}
    path=os.path.abspath(docpath)
    name=path+'x'
    word=Dispatch('Word.Application')
    word.Visible = 0
    doc = word.Documents.Open(path)
    
    if os.path.exists(name):
        os.remove(name)
    doc.SaveAs(name,ftype[dest_ftype], False, "", True, "", False, False, False, False)
    doc.Close()
    #word.Quit
    return
###################
def LibreOffice_change_type(docpath,dest_ftype='docx'):
    ftype=['doc','docx','txt','html','pdf']
    dest_ftype=dest_ftype.lower()
    if dest_ftype not in ftype:
        dest_ftype=input("输入需要转为后的文件类型:")

    output = subprocess.check_output(["soffice","--headless","--invisible","--convert-to","%s"%dest_ftype,docpath,"--outdir",os.path.split(docpath)[0]])
    return
###############################
def FileTypeChange(docpath,dest_ftype='docx'):
    if sys.platform=='win32':
        MSOffice_change_type(docpath,dest_ftype=dest_ftype)
    elif sys.platform=='linux':    
        LibreOffice_change_type(docpath,dest_ftype=dest_ftype)
    return
##########################
def FileTypeChange_dir(pathdir,dest_ftype='docx'):
    for root,ds,fs in os.walk(pathdir):
        for f in fs:
            dd=os.path.splitext(f)
            if dd[1] in ['.doc','.docx']:
                if re.match('^~\w',dd[0]) is None:
                    path=os.path.join(root,f)
                    print(path)
                    FileTypeChange(path,dest_ftype=dest_ftype)
    return

#################################
def Msdoc2pic(path):
     path=os.path.abspath(path)
     ex=os.path.splitext(path)[1]
     if ex == '.doc':
         Doc2Docx(path)
         path=path+'x'
     doc = Document(path)
     for shape in doc.inline_shapes:
        contID=shape._inline.graphic.graphicData.pic.blipFile.blip.embed
        contTp=doc.part.related_parts[contID].content_type
        if not contTp.startswith('image'):
            continue
        imgN=basename(doc.part.related_parts[contID].partname)
        imgD=doc.part.related_parts[contID]._blob
        with open(imgN,'wb') as fp:
            fp.write(imgD)
            print('ok...')

     if ex == '.doc':            
        os.remove(path)
     return
###########################################
def Msdoc2piczip(path):
     path=os.path.abspath(path)
     ex=os.path.splitext(path)
     if ex[1] == '.doc':
         Doc2Docx(path)
         path=path+'x'

     zipp=ex[0]+'.zip'
     if  not os.path.exists(zipp):
        os.rename(path,zipp)
     f=ZipFile(zipp,'r')
     tmp_path='tmp_path'
     if not os.path.exists(tmp_path):
         os.mkdir(tmp_path)

     for ff in f.namelist():
         f.extract(ff,tmp_path)
     f.close()
     if  not os.path.exists(path):
        os.rename(zipp,path)
     pic=os.listdir(os.path.join(tmp_path,'word/media'))
     if ex[1] == '.doc':            
        os.remove(path)
     return
#############################################    
def ReadDocument(path):
    """
    读取doc和docx文件，输入的文件路径。
    """
    ex=os.path.splitext(path)[1]
    if ex == '.doc':
        try:
            doc=_read_doc(path)
        except:
            path=os.path.abspath(path)
            name=path+'x'
            Doc_To_Docx(path)
            doc=_read_docx(name)
        return doc
    if ex =='.docx':
        doc=_read_docx(path)
        return doc

if __name__=="__main__":
    #inpf=sys.argv[1]
    #doc=ReadDocument(inpf)
    pass
