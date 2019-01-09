#!/usr/bin/env python3
# -*-coding:utf-8-*-

import os,sys
import re
from docx import Document
#need to install python-docx
try:
    import win32com
    from win32com.client import Dispatch, constants
except:
    pass

def _read_doc(path):
    """
    读取doc文件,path是doc文件的路径。
    """
    temp='text.txt'
    if sys.platform=='win32':
        word=Dispatch('Word.Application')
        doc = word.Documents.Open(path)
        doc.SaveAs(temp, 4)
        doc.Close()

        f = open(temp,'r')
        content= f.read()
        f.close()
        
        os.remove(temp)
        
        content = content.replace(" ","")
        #content=re.findall(r"[\u4e00-\u9fa5]+",content)
        #content='\n'.join(content)
        return content
    
    elif sys.platform=='linux':
        content = os.popen('catdoc %s'%path)
        #need to install catdoc for linux,for example:aptitude install catdoc
        content = content.read()
        return content

def _read_docx(path):
    """
    读取docx文件,path为docx文件的路径.
    有关文件可以阅读：http://www.cnblogs.com/wrajj/p/4914102.html
    """
    doc=Document(path)
    fullText=[]
    paras=doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    content='\n'.join(fullText)
    return content

def ReadDocument(path):
    """
    读取doc和docx文件，输入的文件路径。
    """
    ex=os.path.splitext(path)[1]
    if ex == '.doc':
        doc=_read_doc(path)
        return doc
    elif ex =='.docx':
        doc=_read_docx(path)
        return doc

if __name__=="__main__":
    inpf=sys.argv[1]
    doc=ReadDocument(inpf)
