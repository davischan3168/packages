#!/usr/bin/env python3
# -*-coding:utf-8-*-

import codecs
import os,sys
import re
from docx import Document
#need to install python-docx
try:
    import win32com
    from win32com.client import Dispatch, constants
except:
    pass

def ReadFile(filePath,encoding):
    with codecs.open(filePath,"r",encoding) as f:
        return f.read()
def WriteFile(filePath,u,encoding):
    with codecs.open(filePath,"w",encoding) as f:
        f.write(u)

def GBK_2_UTF8(src,dst):
    '''
    定义GBK_2_UTF8方法，用于转换文件存储编码
    '''
    content = ReadFile(src,encoding='gbk')
    WriteFile(dst,content,encoding='utf_8')
    
def _read_doc(path):
    """
    读取doc文件,path是doc文件的路径。
    """
    #temp='text.txt'
    path=os.path.abspath(path)
    if sys.platform=='win32':
        
        word=Dispatch('Word.Application')
        word.Visible = 0
        doc = word.Documents.Open(path)
        fullText=[]
        paras=doc.paragraphs
        for p in paras:
            fullText.append(p.Range.Text)
        content='\n'.join(fullText)        
        #doc.SaveAs(temp, 4)
        #doc.Close()

        #f = open(temp,'r')
        #content= f.read()
        #f.close()
        
        #os.remove(temp)
        
        #content = content.replace(" ","")
        #content=re.findall(r"[\u4e00-\u9fa5]+",content)
        #content='\n'.join(content)
        return content
    
    elif sys.platform=='linux':
        content = os.popen('catdoc %s'%path)
        #need to install catdoc for linux,for example:aptitude install catdoc
        content = content.read()
        return content

def _read_docx(path):
    """
    读取docx文件,path为docx文件的路径.
    有关文件可以阅读：http://www.cnblogs.com/wrajj/p/4914102.html
    """
    doc=Document(path)
    fullText=[]
    paras=doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    content='\n'.join(fullText)
    return content

def pydoc(path):
    """
    读取doc和docx文件，输入的文件路径。
    """
    ex=os.path.splitext(path)[1]
    if ex == '.doc':
        doc=_read_doc(path)
        return doc
    elif ex =='.docx':
        doc=_read_docx(path)
        return doc

title=r"""
\documentclass[myxecjk,msize]{gdhsarticle}%{kindle}
\iffalse
\geometry{paperheight=297mm,%
paperwidth=210mm,
left=2.3cm,%
right=2.3cm,%
top=2.1cm,%
bottom=2.1cm,%
headheight=0cm,%
headsep=0cm,
footskip=0cm
}%
\fi
\linespacing{1.52}%
\pagestyle{empty}
%\include{setting}
%\include{pagesign}
\setlength{\unitlength}{1cm}
\parindent=2em
\definecolor{defaultbgcolor-0}{RGB}{199,237,204}%for eye
\pagecolor{defaultbgcolor-0}
%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%以下正文%%%%%%%%%%%%
\begin{document}
"""

kindle=r"""
\documentclass[myxecjk,msize]{gdhsarticle}%{kindle}
\iffalse
\geometry{paperheight=297mm,%
paperwidth=210mm,
left=2.3cm,%
right=2.3cm,%
top=2.1cm,%
bottom=2.1cm,%
headheight=0cm,%
headsep=0cm,
footskip=0cm
}%
\fi
\linespacing{1.52}%
\pagestyle{empty}
%\include{setting}
%\include{pagesign}
\setlength{\unitlength}{1cm}
\parindent=2em
\definecolor{defaultbgcolor-0}{RGB}{199,237,204}%for eye
\pagecolor{defaultbgcolor-0}
%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%以下正文%%%%%%%%%%%%
\begin{document}
"""

end=r"""\end{document}"""
latexs={'article':title,'kindle':kindle}

def _write_file(path,r):
    f=open(path,'a',encoding='utf8')
    f.write(r)
    f.flush()
    return
def _removef(outpath):
    lsd=['.aux','.log','.out','.xdv','.tex']
    base=os.path.split(outpath)
    basename=os.path.splitext(base[1])[0]
    for ex in lsd:
        try:
            mybsn=basename+ex
            path=os.path.join(base[0],mybsn)
            print(path)
            os.remove(path)
        except Exception as e:
            pass
    return
def _Readfile(inpf):
    try:
        con=open(inpf,encoding='utf-8')
        content=con.readlines()
    except:
        con=open(inpf,encoding='gbk')
        content=con.readlines()        
    con.close()
    
    listd=[]
    trec=re.compile(r'^第\w{1,3}篇')
    trec1=re.compile(r'^第\w{1,3}章')    
    for line in content:
        line=line.strip()
        if trec.match(line) is not None:
            line=r'\section{%s}'%line+'\n\n'
            listd.append(line)
        elif trec1.match(line) is not None:
            line=r'\section{%s}'%line+'\n\n'
            listd.append(line)                    
        
        elif len(line)>0:
            if '%' in line:
                line=re.sub(r'%{1,}','\%',line)+'\n\n'
                #print(line)
                #write_file(outpath,line)
                listd.append(line)
            else:
                line=line+'\n\n'
                #print(line)
                #write_file(outpath,line)
                listd.append(line)
    return listd
    
def pylatex(inpf,oupf,mtype='article'):
    """
    inpf:INput Document for example txt docment
    -----------
    oupf: 输出的文件名称，不包含扩展名
    mtype:article,kindle
    """

    outpath=oupf+'.tex'

    if os.path.exists(outpath):
        os.remove(outpath)
    
    with open(outpath,'a') as f:
        f.writelines(latexs[mtype])
        f.flush()

    if os.path.isfile(inpf):
        try:
            listd=_Readfile(inpf)
            conn=''.join(listd)
            #print(conn)
            #if len(conn)>0:
            _write_file(outpath,conn)
            #else:
            #    print('conn is empty.')
            _write_file(outpath,end)
            os.system('xelatex -no-pdf -interaction=nonstopmode %s' %outpath)
            os.system('xelatex -interaction=nonstopmode %s' %outpath)
            _removef(outpath)        
        except Exception as e:
            print(e)
            pass
    elif isinstance(inpf,list):
        try:
            for li in inpf:
                li=li.strip().replace('\n','\n\n')
                _write_file(outpath,conn)
            _write_file(outpath,end)
            os.system('xelatex -no-pdf -interaction=nonstopmode %s' %outpath)
            os.system('xelatex -interaction=nonstopmode %s' %outpath)
            _removef(outpath)
        except Exception as e:
            print(e)
            pass
            
    

    return

if __name__=="__main__":
    '''
    qyx.csv文件使用GBK编码存储，现在将其转为UTF_8存储
    '''
    #src = 'r.txt'
    #dst = 'qyx'
    #GBK_2_UTF8(src,dst)
    #dd=
    #pylatex(sys.argv[1],sys.argv[2],'article')
    pass
