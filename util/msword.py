#!/usr/bin/env python3
# -*-coding:utf-8-*-

import os
import sys
import re
#import pickle
#import codecs
#import string
#import shutil
from docx import Document
import subprocess
#need to install python-docx
try:
    import win32com
    from win32com.client import Dispatch, constants
except:
    pass

def _read_doc(path):
    """
    读取doc文件,path是doc文件的路径。
    
    temp='text.txt'
    path=os.path.abspath(path)
    if sys.platform=='win32':
        word=Dispatch('Word.Application')
        doc = word.Documents.Open(path)
        doc.SaveAs(temp, 4)
        doc.Close()

        f = open(temp,'r')
        content= f.read()
        f.close()
        
        os.remove(temp)
        
        content = content.replace(" ","")
        #content=re.findall(r"[\u4e00-\u9fa5]+",content)
        #content='\n'.join(content)
    """
    path=os.path.abspath(path)
    if sys.platform=='win32':
        
        word=Dispatch('Word.Application')
        word.Visible = 0
        doc = word.Documents.Open(path)
        fullText=[]
        paras=doc.paragraphs
        for p in paras:
            fullText.append(p.Range.Text)
        content='\n'.join(fullText)          
        return content
    
    elif sys.platform=='linux':
        try:
            content = os.popen('catdoc %s'%path)
            #need to install catdoc for linux,for example:aptitude install catdoc
            content = content.read()
        except:
            content = subprocess.check_output(['antiword',path])
        return content

def _read_docx(path):
    """
    读取docx文件,path为docx文件的路径.
    有关文件可以阅读：http://www.cnblogs.com/wrajj/p/4914102.html
    """
    path=os.path.abspath(path)
    doc=Document(path)
    fullText=[]
    paras=doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    content='\n'.join(fullText)
    return content

def Doc2Docx(path):
    
    path=os.path.abspath(path)
    name=path+'x'
    if sys.platform=='win32':
        word=Dispatch('Word.Application')
        word.Visible = 0
        doc = word.Documents.Open(path)
        
        if os.path.exists(name):
            os.remove(name)
        doc.SaveAs(name,12, False, "", True, "", False, False, False, False)
        doc.Close()
        #word.Quit
    elif sys.platform=='linux':
        output = subprocess.check_output(["soffice","--headless","--invisible","--convert-to","docx",path,"--outdir",os.path.split(path)[0]])

    return

def ReadDocument(path):
    """
    读取doc和docx文件，输入的文件路径。
    """
    ex=os.path.splitext(path)[1]
    if ex == '.doc':
        try:
            doc=_read_doc(path)
        except:
            path=os.path.abspath(path)
            name=path+'x'
            Doc_To_Docx(path)
            doc=_read_docx(name)
        return doc
    elif ex =='.docx':
        doc=_read_docx(path)
        return doc

if __name__=="__main__":
    #inpf=sys.argv[1]
    #doc=ReadDocument(inpf)
    pass
